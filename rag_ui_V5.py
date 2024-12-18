import streamlit as st
import logging
import time
import os
import ollama
import json
import numpy as np
import ollama
from numpy.linalg import norm
from ollama import chat

import PyPDF2
from docx import Document
import pandas as pd

from pymongo import MongoClient

# from llama_index.core import GPTVectorStoreIndex, Document
# from llama_index.core.llms import ChatMessage
# from llama_index.llms.ollama import Ollama


logging.basicConfig(level=logging.INFO)

# Initialize chat history in session state if not already present
if 'messages' not in st.session_state:
    st.session_state.messages = []


# Functions to read different file types
def parse_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.readlines()

def parse_pdf(file_path):
    paragraphs = []
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            paragraphs.append(page.extract_text())
    return paragraphs

def parse_docx(file_path):
    doc = Document(file_path)
    return [para.text for para in doc.paragraphs if para.text]

# Function to stream chat response based on selected model
def stream_chat(model, messages):
    try:
        # Initialize the language model with a timeout
        # llm = Ollama(model=model, request_timeout=120.0) 
        # Stream chat responses from the model
        # resp = llm.stream_chat(messages)
        response = ""
        response_placeholder = st.empty()
        # Append each piece of the response to the output

        for part in chat(model, messages=messages, stream=True):
        #     print(part['message']['content'], end='', flush=True)
        # for r in resp:
            response += part['message']['content']
            response_placeholder.write(response)
        # Log the interaction details
        logging.info(f"Model: {model}, Messages: {messages}, Response: {response}")
        return response
    except Exception as e:
        # Log and re-raise any errors that occur
        logging.error(f"Error during streaming: {str(e)}")
        raise e


def assistant_mode():
    st.title("Chat with LLMs Assistent") 
    # Sidebar for model selection
    model = st.selectbox("Choose a model", ["llama3.1", "phi3", "mistral"])
    logging.info(f"Model selected: {model}")

    # Prompt for user input and save to chat history
    if prompt := st.chat_input("Your question"):
        st.session_state.messages.append({"role": "user", "content": prompt})
        logging.info(f"User input: {prompt}")

        # Display the user's query
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.write(message["content"])

        # Generate a new response if the last message is not from the assistant
        if st.session_state.messages[-1]["role"] != "assistant":
            with st.chat_message("assistant"):
                start_time = time.time()  # Start timing the response generation
                logging.info("Generating response")

                with st.spinner("Writing..."):
                    try:
                        # Prepare messages for the LLM and stream the response
                        messages = [{"role": msg["role"], "content": msg["content"]} for msg in st.session_state.messages]
                        response_message = stream_chat(model, messages)
                        duration = time.time() - start_time  # Calculate the duration
                        response_message_with_duration = f"{response_message}\n\nDuration: {duration:.2f} seconds"
                        st.session_state.messages.append({"role": "assistant", "content": response_message_with_duration})
                        st.write(f"Duration: {duration:.2f} seconds")
                        logging.info(f"Response: {response_message}, Duration: {duration:.2f} s")

                    except Exception as e:
                        # Handle errors and display an error message
                        st.session_state.messages.append({"role": "assistant", "content": str(e)})
                        st.error("An error occurred while generating the response.")
                        logging.error(f"Error: {str(e)}")


def summarization_mode():
    st.title("Summarize your document") 
    # Sidebar for model selection
    model = st.selectbox("Choose a model", ["llama3.1", "phi3", "mistral"])
    logging.info(f"Model selected: {model}")

    # Prompt for user input and save to chat history
    if prompt := st.chat_input("Your question"):
        st.session_state.messages.append({"role": "user", "content": prompt})
        logging.info(f"User input: {prompt}")

        # Display the user's query
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.write(message["content"])

        # Generate a new response if the last message is not from the assistant
        if st.session_state.messages[-1]["role"] != "assistant":
            with st.chat_message("assistant"):
                start_time = time.time()  # Start timing the response generation
                logging.info("Generating response")

                with st.spinner("Writing..."):
                    try:
                        # Prepare messages for the LLM and stream the response
                        messages = [ChatMessage(role=msg["role"], content=msg["content"]) for msg in st.session_state.messages]
                        response_message = stream_chat(model, messages)
                        duration = time.time() - start_time  # Calculate the duration
                        response_message_with_duration = f"{response_message}\n\nDuration: {duration:.2f} seconds"
                        st.session_state.messages.append({"role": "assistant", "content": response_message_with_duration})
                        st.write(f"Duration: {duration:.2f} seconds")
                        logging.info(f"Response: {response_message}, Duration: {duration:.2f} s")

                    except Exception as e:
                        # Handle errors and display an error message
                        st.session_state.messages.append({"role": "assistant", "content": str(e)})
                        st.error("An error occurred while generating the response.")
                        logging.error(f"Error: {str(e)}")
def chatbot_mode():
    # Membaca file TXT
    def read_txt(filename):
        with open(filename, "r", encoding="utf-8-sig") as f:
            return f.read()

    # Membaca file PDF
    def read_pdf(filename):
        with open(filename, "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text

    # Membaca file DOCX
    def read_docx(filename):
        doc = Document(filename)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        return text

    # Membaca file CSV
    # Membaca file CSV tanpa header dan mengisi NaN
    def read_csv(filename):
        df = pd.read_csv(filename, header=None, delimiter=';', skip_blank_lines=True)
        df = df.fillna("0")  # Replace NaN with "0"
        text = ""
        
        # Loop through the rows
        for idx, row in df.iterrows():
            if idx >= 1:  # Ensure that idx is an integer
                text += f"'{idx}. " + ", ".join(str(value) for value in row.values) + "\n"
            else:
                # If it's the first row (header row), don't add an index number
                text += "Ini adalah data csv dengan delimter (,)\n" + ", ".join(str(value) for value in row.values) + "\n"
        
        return text

    # Membaca file Excel (.xlsx)
    def read_xlsx(filename):
        df = pd.read_excel(filename)
        text = ""
        for index, row in df.iterrows():
            text += " ".join(str(value) for value in row.values) + "\n"
        return text

    # Mendapatkan paragraf dari semua file
    def parse_file(filename):
        if filename.endswith(".txt"):
            content = read_txt(filename)
        elif filename.endswith(".pdf"):
            content = read_pdf(filename)
        elif filename.endswith(".docx"):
            content = read_docx(filename)
        elif filename.endswith(".csv"):
            content = read_csv(filename)
        elif filename.endswith(".xlsx"):  # Menambahkan pengecekan untuk file Excel
            content = read_xlsx(filename)
        else:
            raise ValueError(f"Unsupported file type: {filename}")
        
        paragraphs = []
        buffer = []
        for line in content.splitlines():
            line = line.strip()
            if line:
                buffer.append(line)
            elif len(buffer):
                paragraphs.append(" ".join(buffer))
                buffer = []
        if len(buffer):
            paragraphs.append(" ".join(buffer))
        return paragraphs

    # Simpan embedding ke file
    def save_embeddings(filename, embeddings):
        if not os.path.exists("embeddings"):
            os.makedirs("embeddings")
        with open(f"embeddings/{filename}.json", "w") as f:
            json.dump(embeddings, f)

    # Memuat embedding dari file
    def load_embeddings(filename):
        if not os.path.exists(f"embeddings/{filename}.json"):
            return False
        with open(f"embeddings/{filename}.json", "r") as f:
            return json.load(f)

    # Mendapatkan embedding
    def get_embeddings(filename, modelname, chunks):
        if (embeddings := load_embeddings(filename)) is not False:
            return embeddings
        embeddings = [
            ollama.embeddings(model=modelname, prompt=chunk)["embedding"]
            for chunk in chunks
        ]
        save_embeddings(filename, embeddings)
        return embeddings

    # Cosine similarity untuk menemukan kemiripan
    def find_most_similar(needle, haystack):
        needle_norm = norm(needle)
        similarity_scores = [
            np.dot(needle, item) / (needle_norm * norm(item)) for item in haystack
        ]
        return sorted(zip(similarity_scores, range(len(haystack))), reverse=True)

    
    st.title("Chat with LLMs RAG")


    # Load subfolders in the `data` directory
    data_root = "data/"
    subfolders = [f for f in os.listdir(data_root) if os.path.isdir(os.path.join(data_root, f))]
    selected_folder = st.selectbox("Pilih folder dataset:", subfolders) 
    # selected_folder = "HR"

    data_folder = os.path.join(data_root, selected_folder)
    all_paragraphs = []

    for file in os.listdir(data_folder):
        file_path = os.path.join(data_folder, file)
        if file.lower().endswith((".txt", ".pdf", ".docx")):
            if file.endswith(".txt"):
                paragraphs = parse_txt(file_path)
            elif file.endswith(".pdf"):
                paragraphs = parse_pdf(file_path)
            elif file.endswith(".docx"):
                paragraphs = parse_docx(file_path)
            all_paragraphs.extend(paragraphs)

    folder_name = os.path.basename(data_folder)
    embeddings_filename = f"data_embeddings_{folder_name}"
    embeddings = get_embeddings(embeddings_filename, "nomic-embed-text", all_paragraphs)

    # Sidebar for model selection
    SYSTEM_PROMPT = """You are an assistant that answers questions only in Bahasa Indonesia. 
    Your answers must be based solely on the provided context extracted from the documents. 
    If the answer cannot be determined from the context, respond with \"Maaf, saya tidak tahu.\" 
    Do not include any information outside of the given context, and strictly reply in Bahasa Indonesia.

    Context:
    """
    model = "llama3.1"
    logging.info(f"Model selected: {model}")
    


    # Prompt for user input and save to chat history
    if prompt := st.chat_input("Your question"):
        st.session_state.messages.append({"role": "user", "content": prompt})

        prompt_embedding = ollama.embeddings(model="nomic-embed-text", prompt=prompt)["embedding"]
        most_similar_chunks = find_most_similar(prompt_embedding, embeddings)[:5]
        context = "\n".join(all_paragraphs[item[1]] for item in most_similar_chunks)
        st.session_state.messages.append({"role": "system", "content": SYSTEM_PROMPT + context})
        
        
        logging.info(f"User input: {prompt}")

        # Display the user's query
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.write(message["content"])

        # Generate a new response if the last message is not from the assistant
        if st.session_state.messages[-1]["role"] != "assistant":
            with st.chat_message("assistant"):
                start_time = time.time()  # Start timing the response generation
                logging.info("Generating response")

                with st.spinner("Writing..."):
                    try:
                        # Prepare messages for the LLM and stream the response
                        messages = [{"role": msg["role"], "content": msg["content"]} for msg in st.session_state.messages]
                        response_message = stream_chat(model, messages)
                        duration = time.time() - start_time  # Calculate the duration
                        response_message_with_duration = f"{response_message}\n\nDuration: {duration:.2f} seconds"
                        st.session_state.messages.append({"role": "assistant", "content": response_message_with_duration})
                        st.write(f"Duration: {duration:.2f} seconds")
                        logging.info(f"Response: {response_message}, Duration: {duration:.2f} s")

                    except Exception as e:
                        # Handle errors and display an error message
                        st.session_state.messages.append({"role": "assistant", "content": str(e)})
                        st.error("An error occurred while generating the response.")
                        logging.error(f"Error: {str(e)}")

def main():
     # Set the title of the Streamlit app
    logging.info("App started")  # Log that the app has started

    SYSTEM_PROMPT = """You are an assistant that answers questions only in Bahasa Indonesia. 
    Your answers must be based solely on the provided context extracted from the documents. 
    If the answer cannot be determined from the context, respond with \"Maaf, saya tidak tahu.\" 
    Do not include any information outside of the given context, and strictly reply in Bahasa Indonesia.

    Context:
    """

    # Sidebar settings
    st.sidebar.title("Pengaturan")
    feature = st.sidebar.radio("Pilih fitur:", ["Chatbot", "Summarization", "Assistant"])

    if feature == "Chatbot":
        # Select the model to use for chat
        chatbot_mode()
    elif feature == "Summarization":
        # Select the model to use for summarization
        summarization_mode()
    elif feature == "Assistant":
        # Select the model to use for assistant
        assistant_mode()
    

if __name__ == "__main__":
    main()